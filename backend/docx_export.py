"""
Builds a clean, ATS-friendly .docx resume from resume JSON data.

ATS-safe layout rules followed:
- Single column, no tables/text boxes/graphics
- Standard heading styles
- No headers/footers containing critical info
- Simple bullet lists
"""
import io

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models import ResumeExportRequest


def _add_heading(document: Document, text: str):
    heading = document.add_heading(text, level=2)
    for run in heading.runs:
        run.font.size = Pt(13)


def _add_contact_line(document: Document, req: ResumeExportRequest):
    parts = []
    if req.contact.email:
        parts.append(req.contact.email)
    if req.contact.phone:
        parts.append(req.contact.phone)
    if req.contact.location:
        parts.append(req.contact.location)
    if req.contact.linkedin:
        parts.append(req.contact.linkedin)
    if req.contact.github:
        parts.append(req.contact.github)

    p = document.add_paragraph(" | ".join(parts))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)


def _add_skill_line(document: Document, label: str, items):
    if not items:
        return
    p = document.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(", ".join(items))


def build_resume_docx(req: ResumeExportRequest) -> io.BytesIO:
    document = Document()

    # Base font
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Name (title)
    title = document.add_heading(req.name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if req.target_role:
        role_p = document.add_paragraph(req.target_role)
        role_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in role_p.runs:
            run.italic = True

    _add_contact_line(document, req)

    # Professional Summary
    if req.summary:
        _add_heading(document, "Professional Summary")
        document.add_paragraph(req.summary)

    # Skills
    skills = req.skills
    if any([skills.automation, skills.api, skills.ci_cd, skills.languages, skills.tools]):
        _add_heading(document, "Skills")
        _add_skill_line(document, "Automation", skills.automation)
        _add_skill_line(document, "API Testing", skills.api)
        _add_skill_line(document, "CI/CD", skills.ci_cd)
        _add_skill_line(document, "Languages", skills.languages)
        _add_skill_line(document, "Tools", skills.tools)

    # Experience
    if req.experience:
        _add_heading(document, "Experience")
        for exp in req.experience:
            p = document.add_paragraph()
            run = p.add_run(f"{exp.title} — {exp.company}")
            run.bold = True
            if exp.duration:
                p.add_run(f"  ({exp.duration})")

            if exp.tech_stack:
                tech_p = document.add_paragraph()
                tech_run = tech_p.add_run("Tech stack: ")
                tech_run.italic = True
                tech_p.add_run(", ".join(exp.tech_stack))

            if exp.responsibilities:
                for line in exp.responsibilities.split("\n"):
                    line = line.strip("-• ").strip()
                    if line:
                        document.add_paragraph(line, style="List Bullet")

            if exp.metrics:
                metrics_p = document.add_paragraph()
                metrics_run = metrics_p.add_run("Impact: ")
                metrics_run.bold = True
                metrics_p.add_run(exp.metrics)

    # Projects
    if req.projects:
        _add_heading(document, "Projects")
        for proj in req.projects:
            p = document.add_paragraph()
            run = p.add_run(proj.name)
            run.bold = True
            if proj.stack:
                p.add_run(f"  ({proj.stack})")
            if proj.description:
                document.add_paragraph(proj.description)
            if proj.impact:
                impact_p = document.add_paragraph()
                impact_run = impact_p.add_run("Impact: ")
                impact_run.bold = True
                impact_p.add_run(proj.impact)

    # Education
    if req.education:
        _add_heading(document, "Education")
        for edu in req.education:
            p = document.add_paragraph()
            run = p.add_run(f"{edu.degree}, {edu.institution}")
            run.bold = True
            if edu.year:
                p.add_run(f"  ({edu.year})")

    # Certifications
    if req.certifications:
        _add_heading(document, "Certifications")
        for cert in req.certifications:
            document.add_paragraph(cert, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
